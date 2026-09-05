"""Generate approval-note DOCX deliverables."""

from pathlib import Path
from tempfile import NamedTemporaryFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.models.approval_workflow import ApprovalWorkflowResult


class ApprovalNoteGenerationError(RuntimeError):
    """Raised when an approval note cannot be generated."""


class ApprovalNoteGenerator:
    """Generate a deterministic approval note as a DOCX file."""

    def generate(
        self,
        result: ApprovalWorkflowResult,
        output_path: Path,
    ) -> Path:
        if not isinstance(result, ApprovalWorkflowResult):
            raise ApprovalNoteGenerationError(
                "result must be an ApprovalWorkflowResult"
            )

        if not output_path:
            raise ApprovalNoteGenerationError("output_path is required")

        output_path = Path(output_path)

        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)

            document = Document()

            title = document.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run("Inspection Approval Note")
            run.bold = True
            run.font.size = Pt(18)

            document.add_paragraph()

            workflow_paragraph = document.add_paragraph()
            workflow_paragraph.add_run("Workflow ID: ").bold = True
            workflow_paragraph.add_run(str(result.workflow_id))

            decision_paragraph = document.add_paragraph()
            decision_paragraph.add_run("Decision: ").bold = True
            decision_paragraph.add_run(result.decision.value.upper())

            document.add_heading("Summary", level=2)
            document.add_paragraph(result.summary)

            document.add_heading("Supporting Evidence", level=2)

            if result.supporting_evidence:
                for evidence in result.supporting_evidence:
                    document.add_paragraph(evidence, style="List Bullet")
            else:
                document.add_paragraph("No supporting evidence provided.")

            with NamedTemporaryFile(
                suffix=".docx",
                dir=output_path.parent,
                delete=False,
            ) as temporary_file:
                temporary_path = Path(temporary_file.name)

            try:
                document.save(temporary_path)
                temporary_path.replace(output_path)
            finally:
                if temporary_path.exists():
                    temporary_path.unlink()

        except (OSError, ValueError) as exc:
            raise ApprovalNoteGenerationError(
                f"Failed to generate approval note: {exc}"
            ) from exc

        return output_path
